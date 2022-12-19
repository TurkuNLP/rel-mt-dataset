import docx
import sys
import tqdm

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

    
for fname in sys.argv[1:]:
    d=docx.Document(docx=fname)
    cut=len(d.paragraphs)//2
    newd=docx.Document()
    for idx,p in tqdm.tqdm(enumerate(d.paragraphs)):
        if p.runs and p.runs[0].bold and idx>cut:
            cut=999999999999999999999999999999999999
            newd.save(fname.replace(".docx","_part1.docx").replace(" ","_"))
            newd=docx.Document()
        get_para_data(newd,p)
    else:
        newd.save(fname.replace(".docx","_part2.docx").replace(" ","_"))
            
    
