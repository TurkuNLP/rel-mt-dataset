import docx
import json
import sys

def index_fonts():
    palette=[]
    with open("palette.txt") as f:
        for line in f:
            line=line.strip()
            if not line.startswith("#"):
                continue
            palette.append(line)
    #print(palette)
    assert len(palette)==len(set(palette))

    rgb_colors=[]
    for colorspec in palette:
        color=docx.shared.RGBColor.from_string(colorspec[1:].upper())
        rgb_colors.append(color)

    return rgb_colors

rgb_colors=index_fonts()


def emit_doc(doc,outdoc,rgb_colors):
    #Do we have overlaps? God I hope not!
    entitycounts=[0 for _ in doc["sentence"]]
    for b,e,_ in doc["ner"]:
        for i in range(b,e+1):
            entitycounts[i]+=1
    for c in entitycounts:
        if c>1: #o ou
            assert False
    ner_spans=[]
    for ner_idx,(t1,t2,_) in enumerate(doc["ner"]):
        ner_spans.append((ner_idx,doc["token2charspan"][t1][0],doc["token2charspan"][t2][1]))
    #these are ner_spans (idx,b,e)
    final_spans=[]
    for (idx,b,e) in ner_spans:
        #is there any text before me?
        if final_spans:
            prev_idx=final_spans[-1][2]
        else:
            prev_idx=0
        if prev_idx<b:
            final_spans.append((-1,prev_idx,b))
        final_spans.append((idx,b,e))
    else:
        if final_spans:
            prev_idx=final_spans[-1][2]
        else:
            prev_idx=0
        if prev_idx<len(doc["sentence-detokenized"]):
            final_spans.append((-1,prev_idx,len(doc["sentence-detokenized"])))

    para=outdoc.add_paragraph("")
    for color,b,e in final_spans:
        run=para.add_run(doc["sentence-detokenized"][b:e])
        if color>=0:
            run.font.color.rgb=rgb_colors[color]
    
    assert para.text==doc["sentence-detokenized"], (para.text,doc["sentence-detokenized"])



outdoc=docx.Document()


for line in sys.stdin:
    doc=json.loads(line)
    outdoc.add_paragraph("").add_run("Page").bold=True
    emit_doc(doc,outdoc,rgb_colors)

outdoc.save("out2.docx")

    
